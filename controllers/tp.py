# from docx import Document
# from docx.shared import Pt, RGBColor
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
# from docx2pdf import convert
# import smtplib
# import ssl
# from email.message import EmailMessage
# import sys
# import pytesseract
# import cv2

# image_path = sys.argv[1]
# image = cv2.imread(image_path)

# height, width, _ = image.shape

# upper_left = (0, 0)
# upper_right = (width, height // 2)
# lower_left = (0, height // 2)
# lower_right = (width, height)

# upper_half = image[upper_left[1]:upper_right[1], upper_left[0]:upper_right[0]]

# lower_half = image[lower_left[1]:lower_right[1], lower_left[0]:lower_right[0]]

# upper_text = pytesseract.image_to_string(upper_half)

# lower_text = pytesseract.image_to_string(lower_half)

# symptoms = upper_text.strip()
# medicines = lower_text.strip()

# symptoms_text = "Symptoms:\n" + symptoms
# medicines_text = "Medicines:\n" + medicines

# import json

# data = {
#     "symptoms": symptoms,
#     "medicines": medicines
# }

# json_path = 'temp.json'

# with open(json_path, 'w') as json_file:
#     json.dump(data, json_file)

# print("Symptoms and medicines data saved to:", json_path)

# doc = Document('aa2.docx')

# symptoms_paragraph = doc.add_paragraph(symptoms_text)

# medicines_paragraph = doc.add_paragraph(medicines_text)

# doc_path = 'final.docx'
# doc.save(doc_path)

# pdf_path = 'final.pdf'
# convert(doc_path, pdf_path)

# def send_email():
#     email_sender = 'mananparekh1755@gmail.com'
#     email_password = "wlrp jpmu ckbg bmlm"
#     email_receiver = 'mananparekh.co22d2@scet.ac.in'

#     subject = 'Check out my new mail'
#     body = "I've just sent you an email with an attachment."

#     em = EmailMessage()
#     em['From'] = email_sender
#     em['To'] = email_receiver
#     em['Subject'] = subject
#     em.set_content(body)

#     with open(pdf_path, 'rb') as attachment_file:
#         file_data = attachment_file.read()
#         file_name = attachment_file.name.split("/")[-1]

#         em.add_attachment(file_data, maintype='application', subtype='octet-stream', filename=file_name)

#     context = ssl.create_default_context()

#     with smtplib.SMTP_SSL('smtp.gmail.com', 465, context=context) as smtp:
#         smtp.login(email_sender, email_password)
#         smtp.send_message(em)

# send_email()

# print(symptoms_text)
# print(medicines_text)

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import smtplib
import ssl
from email.message import EmailMessage
import sys
import pytesseract
import cv2
import json

image_path = sys.argv[1]
image = cv2.imread(image_path)

height, width, _ = image.shape

upper_left = (0, 0)
upper_right = (width, height // 2)
lower_left = (0, height // 2)
lower_right = (width, height)

upper_half = image[upper_left[1]:upper_right[1], upper_left[0]:upper_right[0]]

lower_half = image[lower_left[1]:lower_right[1], lower_left[0]:lower_right[0]]

upper_text = pytesseract.image_to_string(upper_half)

lower_text = pytesseract.image_to_string(lower_half)

symptoms = upper_text.strip()
medicines = lower_text.strip()

symptoms_text = "Symptoms:\n" + symptoms
medicines_text = "Medicines:\n" + medicines

data = {
    "symptoms": symptoms,
    "medicines": medicines
}

json_path = 'temp.json'

with open(json_path, 'w') as json_file:
    json.dump(data, json_file)

print("Symptoms and medicines data saved to:", json_path)

doc = Document('aa2.docx')

symptoms_paragraph = doc.add_paragraph(symptoms_text)

medicines_paragraph = doc.add_paragraph(medicines_text)

doc_path = 'final.docx'
doc.save(doc_path)

pdf_path = 'final.pdf'
convert(doc_path, pdf_path)

def send_email():
    email_sender = 'mananparekh1755@gmail.com'
    email_password = "wlrp jpmu ckbg bmlm"
    email_receiver = 'mananparekh.co22d2@scet.ac.in'

    subject = 'Check out my new mail'
    body = "I've just sent you an email with an attachment."

    em = EmailMessage()
    em['From'] = email_sender
    em['To'] = email_receiver
    em['Subject'] = subject
    em.set_content(body)

    with open(pdf_path, 'rb') as attachment_file:
        file_data = attachment_file.read()

        em.add_attachment(file_data, maintype='application', subtype='octet-stream', filename='final.pdf')

    context = ssl.create_default_context()

    with smtplib.SMTP_SSL('smtp.gmail.com', 465, context=context) as smtp:
        smtp.login(email_sender, email_password)
        smtp.send_message(em)

send_email()

print(symptoms_text)
print(medicines_text)
